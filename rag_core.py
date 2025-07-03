
# --- Standard Library Imports ---
import base64
import hashlib
import io
import os
import tempfile
import logging

# --- Third-Party Imports ---
import docx
import faiss
import numpy as np
import pdfplumber
from dotenv import load_dotenv
from PIL import Image
from openai import AzureOpenAI

# LangChain imports
from langchain.docstore.document import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain_openai import AzureOpenAIEmbeddings

# --- Configuration & Initialization ---
load_dotenv()
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Azure OpenAI credentials
AZURE_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION")
AZURE_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT")
AZURE_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")
COMPLETION_DEPLOYMENT = os.getenv("AZURE_OPENAI_COMPLETION_DEPLOYMENT_NAME")
EMBEDDING_DEPLOYMENT = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT_NAME")

# Early validation
if not all([AZURE_API_VERSION, AZURE_ENDPOINT, AZURE_API_KEY, COMPLETION_DEPLOYMENT, EMBEDDING_DEPLOYMENT]):
    raise ValueError("Missing Azure OpenAI environment variables. Check your .env file.")

# Initialize models
try:
    vision_client = AzureOpenAI(
        api_version=AZURE_API_VERSION,
        azure_endpoint=AZURE_ENDPOINT,
        api_key=AZURE_API_KEY,
    )
    EMBEDDINGS = AzureOpenAIEmbeddings(
        api_version=AZURE_API_VERSION,
        azure_endpoint=AZURE_ENDPOINT,
        api_key=AZURE_API_KEY,
        model="text-embedding-ada-002",
        azure_deployment=EMBEDDING_DEPLOYMENT
    )
except Exception as e:
    logging.error(f"Model initialization failed: {e}")
    raise

# FAISS storage configuration
FAISS_INDEX_PATH = "faiss_index"
os.makedirs(FAISS_INDEX_PATH, exist_ok=True)
TEXT_SPLITTER = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=200)

# --- Language Prompts ---
LANGUAGE_PROMPTS = {
    'English': {
        'ocr': """Extract ALL text from this image exactly as written. Do not modify or interpret.
            Instructions:
            - Only give answer based on provided context.
            - If user question is irrelevant to context. Just write : Information not availaible.""",
        'qa': """Answer the query based ONLY on this context. Respond in English. 
        Instructions:
        - Only give answer based on provided context.
        - If user question is irrelevant to context. Just write : Information not availaible.
Context: {context}
Query: {question}"""
    },
    'Japanese': {
        'ocr': """画像内のテキストをすべてそのまま抽出してください。修正や解釈は不要です。
                指示：
    - 提供されたコンテキストに基づいてのみ回答してください。
    - ユーザーの質問がコンテキストに関連しない場合は、「情報がありません」とだけ記入してください。""",
        'qa': """このコンテキストに基づいて質問に答えてください。日本語で回答してください。
        指示：
- 提供されたコンテキストに基づいてのみ回答してください。
- ユーザーの質問がコンテキストに関連しない場合は、「情報がありません」とだけ記入してください。
コンテキスト: {context}
質問: {question}"""
    },
    'Chinese': {
        'ocr': """准确提取图片中的所有文字，不要修改或解释。
        说明：
- 仅根据提供的上下文给出答案。
- 如果用户问题与上下文无关，请填写：信息不可用。""",
        'qa': """仅根据以下内容回答问题。用中文回答。
        说明：
- 仅根据提供的上下文给出答案。
- 如果用户问题与上下文无关，请填写：信息不可用。
内容: {context}
问题: {question}"""
    }
}


# --- Core Functions ---
def extract_text_from_image(image: Image.Image, language: str) -> str:
    buffered = io.BytesIO()
    image.save(buffered, format="PNG")
    img_base64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
    try:
        response = vision_client.chat.completions.create(
            model=COMPLETION_DEPLOYMENT,
            messages=[{
                "role": "user",
                "content": [
                    {"type": "text", "text": LANGUAGE_PROMPTS[language]['ocr']},
                    {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_base64}"}}
                ]
            }],
            max_tokens=2000
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        logging.warning(f"OCR failed: {str(e)}")
        return ""

def process_pdf(file_bytes: bytes, language: str) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"\n--- Page {i+1} ---\n{page_text}")
            for img in page.images:
                try:
                    img_data = img['stream'].get_data()
                    img_pil = Image.open(io.BytesIO(img_data))
                    ocr_text = extract_text_from_image(img_pil, language)
                    if ocr_text:
                        text_parts.append(f"\n--- Page {i+1} Image ---\n{ocr_text}")
                except Exception as e:
                    logging.warning(f"Image processing failed on page {i+1}: {str(e)}")
                    continue
    return "\n".join(text_parts)

def process_docx(file_bytes: bytes) -> str:
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file_bytes)
        tmp_path = tmp.name
    try:
        doc = docx.Document(tmp_path)
        return "\n".join([p.text for p in doc.paragraphs if p.text])
    finally:
        os.remove(tmp_path)

def process_uploaded_file(file_bytes: bytes, file_type: str, file_name: str, language: str) -> str:
    if file_type == "application/pdf":
        return process_pdf(file_bytes, language)
    elif file_type.startswith("image/"):
        return extract_text_from_image(Image.open(io.BytesIO(file_bytes)), language)
    elif file_name.endswith(".docx"):
        return process_docx(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def ingest_document(file_bytes: bytes, file_type: str, file_name: str, language: str) -> str:
    logging.info(f"Starting ingestion for file: {file_name}")
    doc_hash = hashlib.sha256(file_bytes).hexdigest()
    doc_id = f"doc_{doc_hash}"
    doc_path = os.path.join(FAISS_INDEX_PATH, doc_id)

    if os.path.exists(doc_path):
        logging.info(f"Document {doc_id} already processed. Skipping.")
        return doc_id

    try:
        text = process_uploaded_file(file_bytes, file_type, file_name, language)
        if not text.strip():
            raise ValueError("No text could be extracted from the document.")
        
        docs = TEXT_SPLITTER.split_documents([
            LangchainDocument(page_content=text, metadata={"source": file_name, "language": language})
        ])
        
        vectorstore = FAISS.from_documents(docs, EMBEDDINGS)
        vectorstore.save_local(doc_path)
        logging.info(f"Successfully processed and saved document {doc_id}")
        return doc_id
    except Exception as e:
        logging.error(f"Ingestion failed for {file_name}: {e}")
        raise

def query_document(doc_id: str, question: str, language: str) -> str:
    logging.info(f"Querying document {doc_id} with question: '{question}'")
    doc_path = os.path.join(FAISS_INDEX_PATH, doc_id)
    if not os.path.exists(doc_path):
        raise FileNotFoundError(f"No processed document found with ID: {doc_id}")
        
    try:
        # **FIX APPLIED HERE**: Removed 'allow_dangerous_deserialization' for compatibility
        vectorstore = FAISS.load_local(doc_path, EMBEDDINGS)
        context = "\n".join([doc.page_content for doc in vectorstore.similarity_search(question, k=3)])
        
        system_prompt = LANGUAGE_PROMPTS[language]['qa'].split('Context:')[0].strip()
        user_content = f"Context:\n{context}\n\nQuestion:\n{question}"

        response = vision_client.chat.completions.create(
            model=COMPLETION_DEPLOYMENT,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content}
            ],
            temperature=0.7,
            max_tokens=2000
        )
        answer = response.choices[0].message.content
        logging.info(f"Generated answer for document {doc_id}")
        return answer
    except Exception as e:
        logging.error(f"Query failed for {doc_id}: {e}")
        raise